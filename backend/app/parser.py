import os
import re
from typing import List, Dict, Any
from docx import Document
from pypdf import PdfReader

class QuizParseError(Exception):
    pass

def parse_quiz_text(raw_text: str) -> List[Dict[str, Any]]:
    """
    Parses the custom quiz format:
    +++
    Question?
    ===
    #Correct Answer
    ===
    Incorrect Answer 1
    ===
    Incorrect Answer 2
    +++
    
    Upgrades:
    - Splitting with regex boundary \+{3,} and ={3,} to ignore extra signs or whitespaces.
    - Robust UTF-8 and carriage return cleanup.
    - Support for `#Option`, `# Option` and `Option #`.
    """
    # Normalize carriage returns and newlines
    raw_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Split text by 3 or more '+' boundary (e.g., +++, ++++)
    blocks = re.split(r'\+{3,}', raw_text)
    questions = []
    
    for i, block in enumerate(blocks):
        block = block.strip()
        if not block:
            continue
            
        # Split by 3 or more '=' boundary (e.g., ===, ====)
        parts = [p.strip() for p in re.split(r'={3,}', block) if p.strip()]
        
        if len(parts) < 3:
            # Not a valid question block (must have at least question text and 2 options)
            continue
            
        question_text = parts[0]
        options = parts[1:]
        
        correct_index = -1
        cleaned_options = []
        
        for idx, option in enumerate(options):
            opt_stripped = option.strip()
            # Support both '#Option', '# Option' and 'Option #'
            if opt_stripped.startswith("#") or opt_stripped.endswith("#"):
                if correct_index != -1:
                    raise QuizParseError(
                        f"Xatolik: '{question_text[:50]}...' savolida birdan ortiq to'g'ri javob (#) topildi."
                    )
                correct_index = idx
                # Remove '#' and strip spaces
                if opt_stripped.startswith("#"):
                    cleaned_options.append(opt_stripped[1:].strip())
                else:
                    cleaned_options.append(opt_stripped[:-1].strip())
            else:
                cleaned_options.append(opt_stripped)
                
        if correct_index == -1:
            # Skip blocks that have no correct answer marked (likely intro or explanation text)
            continue
            
        if len(cleaned_options) < 2:
            continue
            
        if len(cleaned_options) > 10:
            raise QuizParseError(
                f"Xatolik: '{question_text[:50]}...' savolida variantlar soni 10 tadan ko'p ({len(cleaned_options)} ta). Telegram testlarida maksimal 10 ta variant bo'lishi mumkin."
            )
            
        questions.append({
            "question_text": question_text,
            "options": cleaned_options,
            "correct_option_index": correct_index
        })
        
    if not questions:
        raise QuizParseError(
            "Faylda hech qanday yaroqli test savoli topilmadi!\n\n"
            "Iltimos, faylingizda quyidagi shartlarga to'liq amal qilinganini tekshiring:\n"
            "1. Har bir savol boshlanishi va tugashida `+++` belgilari bo'lishi shart.\n"
            "2. Savol va variantlar orasida `===` ajratuvchisi bo'lishi shart.\n"
            "3. To'g'ri javob variantining boshida (yoki oxirida) `#` belgisi bo'lishi shart.\n"
            "4. Variantlar soni 2 tadan 10 tagacha bo'lishi shart."
        )
        
    return questions

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from paragraphs and tables inside a DOCX file.
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # 1. Read all normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # 2. Read all tables and cells (in case questions are inside table grids!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        # Standardize boundary formatting inside tables
                        full_text.append(cell.text)
                        
        text = "\n".join(full_text)
        return parse_quiz_text(text)
    except Exception as e:
        if isinstance(e, QuizParseError):
            raise e
        raise QuizParseError(f"DOCX faylini o'qishda xatolik yuz berdi: {str(e)}")

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a PDF file and parses it.
    """
    try:
        reader = PdfReader(file_path)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        text = "\n".join(full_text)
        return parse_quiz_text(text)
    except Exception as e:
        if isinstance(e, QuizParseError):
            raise e
        raise QuizParseError(f"PDF faylini o'qishda xatolik yuz berdi: {str(e)}")

def parse_txt(file_path: str) -> List[Dict[str, Any]]:
    """
    Reads a TXT file and parses it.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return parse_quiz_text(text)
    except Exception as e:
        if isinstance(e, QuizParseError):
            raise e
        raise QuizParseError(f"TXT faylini o'qishda xatolik yuz berdi: {str(e)}")

def parse_quiz_file(file_path: str) -> List[Dict[str, Any]]:
    """
    Routes the file to the appropriate parser based on extension.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".txt", ".doc"]:
        try:
            return parse_txt(file_path)
        except Exception:
            raise QuizParseError("Eski .DOC formatini o'qib bo'lmadi. Iltimos, uni .DOCX yoki .TXT formatiga o'tkazib yuboring.")
    else:
        raise QuizParseError(f"Qo'llab-quvvatlanmaydigan fayl formati: {ext}. Faqat DOCX, PDF, TXT fayllari qabul qilinadi.")
